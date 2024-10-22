import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def copiar_documento_word(arquivo_origem, arquivo_destino):
    # Abrir o documento de origem
    doc_origem = docx.Document(arquivo_origem)

    # Criar um novo documento
    novo_doc = docx.Document()

    # Copiar parágrafos e manter a formatação
    for para in doc_origem.paragraphs:
        # Adicionar um novo parágrafo no novo documento
        novo_para = novo_doc.add_paragraph()

        # Copiar a formatação 
        for run in para.runs:
            novo_run = novo_para.add_run(run.text)
            novo_run.bold = run.bold

            # Configurar a cor da fonte
            if run.font.color and run.font.color.rgb:  
                novo_run.font.color.rgb = run.font.color.rgb
            else:
                novo_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  

            # Copiar tamanho da fonte
            if run.font.size:  
                novo_run.font.size = run.font.size
            else:
                novo_run.font.size = docx.shared.Pt(12) 

            # Copiar nome da fonte
            if run.font.name:  
                novo_run.font.name = run.font.name
            else:
                novo_run.font.name = "Arial"  

        # Copiar o espaçamento do parágrafo
        novo_para.paragraph_format.space_before = para.paragraph_format.space_before
        novo_para.paragraph_format.space_after = para.paragraph_format.space_after
        novo_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing

        # Aplicar o estilo do parágrafo
        novo_para.style = para.style
        
        # Justificar o parágrafo
        novo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Salvar o novo documento
    novo_doc.save(arquivo_destino)
    print(f"Cópia gerada com sucesso: {arquivo_destino}")

arquivo_origem = "python.docx"

pasta_destino = "docs"

# Caminho completo do arquivo de destino dentro da pasta 'docs'
arquivo_destino = os.path.join(pasta_destino, "python_copia.docx")

# Copiar o documento com formatação
copiar_documento_word(arquivo_origem, arquivo_destino)
